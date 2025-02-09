import os
import datetime
from docx import Document
from docx.shared import Pt

def replace_placeholders(paragraph, replacements):
    """Replaces placeholders in a given paragraph while preserving formatting."""
    for placeholder, value in replacements.items():
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)
                    if placeholder in ["{{JOB_POSITION}}", "{{COMPANY_NAME}}"]:
                        run.bold = True
                        run.font.size = Pt(16)

def process_word_document(template_path, output_path, replacements):
    """Reads a Word document, replaces placeholders, and saves a new version while preserving formatting."""
    doc = Document(template_path)
    for para in doc.paragraphs:
        replace_placeholders(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_placeholders(para, replacements)
    doc.save(output_path)

def generate_resume_and_cover(job_position, company_name, specific_job_project, required_it_skills):
    """Generates a customized resume and cover letter based on provided job details."""
    
    # Date formatting
    current_date = datetime.datetime.now().strftime("%B %d, %Y")
    folder_date = datetime.datetime.now().strftime("%m%d%y")
    
    # Output folder setup
    output_folder = os.path.join("output", f"{company_name}_{folder_date}")
    os.makedirs(output_folder, exist_ok=True)
    
    # File paths
    resume_template = os.path.join("static", "resume_template.docx")
    cover_letter_template = os.path.join("static", "cover_letter_template.docx")
    resume_output = os.path.join(output_folder, "henry_pai_resume.docx")
    cover_letter_output = os.path.join(output_folder, "henry_pai_cover_letter.docx")
    
    # Replacements dictionary
    replacements = {
        "{{JOB_POSITION}}": job_position,
        "{{COMPANY_NAME}}": company_name,
        "{{SPECIFIC_JOB_PROJECT}}": specific_job_project,
        "{{IT_SKILLS}}": required_it_skills,
        "{{CURRENT_DATE}}": current_date,
    }
    
    # Process documents
    process_word_document(resume_template, resume_output, replacements)
    process_word_document(cover_letter_template, cover_letter_output, replacements)
    
    print(f"Customized resume saved at: {resume_output}")
    print(f"Customized cover letter saved at: {cover_letter_output}")
    
    return output_folder  # Return path for Tkinter confirmation

# Only runs if this file is executed directly (not when imported)
if __name__ == "__main__":
    generate_resume_and_cover(
        job_position="Data Scientist",
        company_name="Flip",
        specific_job_project="Leveraging data to drive innovation in the e-commerce industry",
        required_it_skills="SQL, Python, R, ML, statistics, etc."
    )
